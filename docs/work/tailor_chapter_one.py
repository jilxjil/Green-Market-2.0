from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "docs" / "work" / "original_chapter_one.docx"
OUTPUT = ROOT / "docs" / "output" / "green_market_chapter_one_tailored.docx"


chapter_one = [
    ("CHAPTER ONE: INTRODUCTION", "Title"),
    ("1.1 Introduction", "Heading 1"),
    (
        "Green Market is a web-based agricultural marketplace and advisory platform designed to connect farmers, buyers, agricultural experts, and administrators within one secure digital ecosystem. The system responds to common challenges in the agricultural value chain, especially limited market access for farmers, fragmented communication between stakeholders, and difficulty in finding reliable agricultural support.",
        "Normal",
    ),
    (
        "The current project is implemented as a full-stack Next.js application with role-based access for sellers, buyers, experts, and administrators. It supports product listing, order management, buyer shopping workflows, expert service publishing, consultation requests, reviews, notifications, messaging, and dashboard-based monitoring. These features allow agricultural stakeholders to interact through a structured platform instead of relying only on informal channels or middlemen.",
        "Normal",
    ),
    (
        "The platform is built around three practical pillars. The first is a digital marketplace where sellers can publish farm products with descriptions, categories, images, prices, units of measure, stock quantities, and product status. The second is an expert service module where agricultural professionals can offer paid consultation services and manage requests from users who need technical guidance. The third is a management and accountability layer that includes user profiles, role-based dashboards, order fulfilment tracking, reviews, notifications, conversations, and administrative oversight.",
        "Normal",
    ),
    (
        "By combining these modules, Green Market aims to improve transparency, reduce communication barriers, support direct agricultural trade, and make expert assistance more accessible. The system therefore contributes to the digital transformation of agriculture by providing a centralized platform for market participation, consultation, and operational management.",
        "Normal",
    ),
    ("1.2 Problem Statement", "Heading 1"),
    (
        "Agriculture remains a major contributor to food security, employment, and economic development in Ghana and other developing economies. Despite its importance, many farmers still face difficulty reaching reliable buyers, presenting their products beyond local markets, and receiving fair value for their produce. Traditional market structures often depend heavily on intermediaries, which can reduce farmer earnings and limit price transparency.",
        "Normal",
    ),
    (
        "Buyers also experience challenges when searching for fresh agricultural products from trusted sellers. The absence of a centralized digital marketplace can make it difficult to compare products, confirm availability, communicate with sellers, and track orders efficiently. This creates uncertainty in pricing, product quality, and supply reliability.",
        "Normal",
    ),
    (
        "Another major challenge is limited access to agricultural expertise. Farmers and other agricultural stakeholders may require guidance on crop management, animal health, pest control, soil improvement, and general farm practices, but professional support is not always easily available. Where advice is available, communication and appointment management may still be informal and difficult to track.",
        "Normal",
    ),
    (
        "There is also a need for better digital coordination within agricultural platforms. User registration, role management, product management, consultation booking, order processing, reviews, messaging, notifications, and administrative monitoring are often handled separately or manually. This fragmentation reduces efficiency and makes it harder to build trust among stakeholders.",
        "Normal",
    ),
    (
        "Therefore, there is a need for an integrated web-based system that brings agricultural marketplace activities, expert consultation services, communication tools, and administrative controls into one platform. Green Market is proposed to address these gaps by improving direct interaction, transparency, accessibility, and digital management across the agricultural value chain.",
        "Normal",
    ),
    ("1.3 Aim of the Project", "Heading 1"),
    (
        "The aim of this project is to design and develop Green Market, a secure and responsive web-based agricultural platform that enables farmers and sellers to market agricultural products directly to buyers while also providing access to expert consultation services and role-based management tools.",
        "Normal",
    ),
    (
        "The project seeks to improve agricultural trade and support services by integrating product listing, order management, expert service booking, messaging, reviews, notifications, user profile management, and administrative oversight into a single digital platform. Through this integration, the system supports more transparent transactions, better communication, and more organized agricultural service delivery.",
        "Normal",
    ),
    ("1.4 Objectives of the Project", "Heading 1"),
    (
        "The objectives of the Green Market project are as follows:",
        "Normal",
    ),
    (
        "To design and develop a user-friendly digital marketplace where farmers and agricultural sellers can create, update, and manage product listings with relevant product details, images, prices, units of measure, stock quantities, and availability status.",
        "List Bullet",
    ),
    (
        "To provide buyers with a convenient platform for browsing agricultural products, managing cart activities, placing orders, and tracking purchase-related information.",
        "List Bullet",
    ),
    (
        "To implement role-based access control for buyers, sellers, experts, and administrators so that each user category can access the functions relevant to its responsibilities.",
        "List Bullet",
    ),
    (
        "To develop an expert services module that allows agricultural experts to publish services, receive consultation requests, manage appointments, and provide guidance to platform users.",
        "List Bullet",
    ),
    (
        "To support communication and trust through conversations, notifications, reviews, seller verification status, and administrative monitoring.",
        "List Bullet",
    ),
    (
        "To build a secure, scalable, and responsive application using modern web technologies including Next.js, TypeScript, PostgreSQL, Drizzle ORM, Better Auth, Tailwind CSS, and Supabase storage support.",
        "List Bullet",
    ),
    ("1.5 Core Features of the System", "Heading 1"),
    ("1.5.1 Digital Marketplace", "Heading 2"),
    (
        "The marketplace is the central trading component of Green Market. It allows sellers to list agricultural products with details such as title, description, category, price, unit of measure, stock quantity, image, and active status. Buyers can access the marketplace to view available products and begin purchase activities through the cart and order flow.",
        "Normal",
    ),
    (
        "The system also supports order records, fulfilment status, shipping address information, and tracking number fields. These features help organize transactions and provide a clearer record of buyer and seller activities.",
        "Normal",
    ),
    ("1.5.2 Expert Services and Consultation Requests", "Heading 2"),
    (
        "The expert services module allows agricultural professionals to create service listings with titles, descriptions, prices, and consultation durations. Users can request consultations, include messages, and schedule sessions. The system also supports consultation status, meeting provider, meeting link, and meeting notes, making expert interactions easier to manage.",
        "Normal",
    ),
    (
        "This feature helps bridge the knowledge gap in agriculture by making professional guidance more accessible through the platform. It also creates a structured way for experts to manage their services and consultation history.",
        "Normal",
    ),
    ("1.5.3 User Profiles and Role-Based Dashboards", "Heading 2"),
    (
        "Green Market provides different profile structures for sellers, buyers, and experts. Sellers can record farm-related information such as farm name, location, and verification status. Buyers can provide business information, while experts can indicate their expertise and years of experience.",
        "Normal",
    ),
    (
        "The system includes dashboards for sellers, buyers, experts, and administrators. These dashboards organize user-specific activities such as product management, order monitoring, consultation management, analytics pages, profile setup, and system administration.",
        "Normal",
    ),
    ("1.5.4 Communication, Reviews, and Notifications", "Heading 2"),
    (
        "To promote trust and collaboration, the platform includes conversations, messages, notifications, and reviews. These features support direct communication among users, help users receive timely updates, and allow participants to provide feedback on marketplace interactions.",
        "Normal",
    ),
    ("1.5.5 Administration and System Monitoring", "Heading 2"),
    (
        "Administrative functions allow authorized users to monitor platform activity, manage users, review sellers, and oversee products. This improves platform governance and helps maintain reliability, accountability, and order within the system.",
        "Normal",
    ),
    ("1.6 Target Users", "Heading 1"),
    (
        "The Green Market platform is designed for the following major user groups:",
        "Normal",
    ),
    (
        "Farmers and sellers: These users list agricultural products, manage stock information, monitor orders, update seller profiles, and interact with buyers through the platform.",
        "List Bullet",
    ),
    (
        "Buyers: Buyers browse products, manage cart activities, place orders, track purchases, and review their marketplace experience.",
        "List Bullet",
    ),
    (
        "Agricultural experts: Experts publish consultation services, receive requests, manage schedules, and provide professional agricultural advice.",
        "List Bullet",
    ),
    (
        "Administrators: Administrators monitor users, sellers, products, and general system activity to support secure and reliable platform operation.",
        "List Bullet",
    ),
    ("1.7 Value Proposition", "Heading 1"),
    (
        "Green Market creates value by providing a centralized platform for agricultural trade, consultation, and communication. For farmers and sellers, the system increases product visibility and supports direct access to buyers. This can reduce overdependence on intermediaries and improve opportunities for fairer pricing.",
        "Normal",
    ),
    (
        "For buyers, the platform improves access to agricultural products by allowing them to view listed items, compare available options, and place orders through a structured digital process. This supports convenience, transparency, and better supply discovery.",
        "Normal",
    ),
    (
        "For agricultural experts, Green Market provides a digital channel for offering professional services, managing consultation requests, and reaching users who need guidance. For administrators, the system provides oversight tools that help maintain trust, security, and proper platform governance.",
        "Normal",
    ),
    (
        "Overall, the platform improves market access, expert accessibility, communication, accountability, and digital organization within the agricultural value chain.",
        "Normal",
    ),
    ("1.8 Scope of the System", "Heading 1"),
    (
        "The scope of this project covers the design and development of a web-based agricultural platform with marketplace, expert consultation, profile management, communication, review, notification, order management, and administrative features. The system is implemented as a responsive full-stack web application using Next.js, TypeScript, PostgreSQL, Drizzle ORM, Better Auth, Tailwind CSS, and Supabase storage support for uploaded images.",
        "Normal",
    ),
    (
        "The system includes authentication, session management, and role-based access control for buyers, sellers, experts, and administrators. It also includes seller product management, buyer shopping workflows, expert service management, consultation request handling, messaging, notifications, reviews, and order fulfilment tracking.",
        "Normal",
    ),
    (
        "The project does not cover physical logistics such as transportation, warehousing, delivery fleet management, or direct farm operations. It also does not claim to fully automate agricultural decision-making. Instead, its focus is on digital market coordination, expert access, transaction organization, and stakeholder communication.",
        "Normal",
    ),
    ("1.9 Project Deliverables", "Heading 1"),
    (
        "Upon completion, the project will produce the following deliverables:",
        "Normal",
    ),
    (
        "A functional Green Market web application with marketplace, expert service, consultation, order, messaging, review, notification, dashboard, and administrative modules.",
        "List Bullet",
    ),
    (
        "The complete source code of the system, including frontend pages, backend API routes, database schemas, migrations, services, validation logic, and supporting scripts.",
        "List Bullet",
    ),
    (
        "A PostgreSQL database design managed through Drizzle ORM, covering users, profiles, products, expert services, consultation requests, orders, order items, reviews, conversations, messages, notifications, and administrators.",
        "List Bullet",
    ),
    (
        "Project documentation explaining the problem, objectives, methodology, system design, implementation, testing, and conclusions.",
        "List Bullet",
    ),
    (
        "A user guide or demonstration material showing how buyers, sellers, experts, and administrators interact with the platform.",
        "List Bullet",
    ),
    ("1.10 Conclusion of Chapter", "Heading 1"),
    (
        "This chapter has introduced the Green Market project by presenting its background, problem statement, aim, objectives, core features, target users, value proposition, scope, and expected deliverables. The chapter establishes the need for a centralized agricultural platform that supports direct marketplace interaction, expert consultation, structured communication, and role-based system management.",
        "Normal",
    ),
    (
        "The proposed system is relevant because it addresses practical challenges faced by agricultural stakeholders, including limited market access, fragmented communication, difficulty in accessing expert advice, and weak digital coordination. The next chapters will review related systems, explain the methodology, and present the design and implementation of the Green Market platform.",
        "Normal",
    ),
]


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def main():
    doc = Document(INPUT)
    chapter_two_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().upper().startswith("CHAPTER TWO"):
            chapter_two_index = index
            break

    if chapter_two_index is None:
        raise RuntimeError("Could not find CHAPTER TWO marker.")

    chapter_two_para = doc.paragraphs[chapter_two_index]
    for paragraph in list(doc.paragraphs[:chapter_two_index]):
        delete_paragraph(paragraph)

    for text, style in chapter_one:
        paragraph = chapter_two_para.insert_paragraph_before(text)
        try:
            paragraph.style = style
        except KeyError:
            paragraph.style = "Normal"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
